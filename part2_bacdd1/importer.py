from docx import Document as DocxDocument
from part2_bacdd1.models import Question

class QuestionImporter:
    def __init__(self, db):
        self.db = db

    def import_from_docx(self, file_path):
        doc = DocxDocument(file_path)

        # Get subject from the first few paragraphs
        subject = None
        for para in doc.paragraphs:
            text = para.text.strip()
            if text.startswith("Subject:"):
                subject = text.replace("Subject:", "").strip()
                break

        if not subject:
            print("[ERROR] Subject not found in document.")
            return False

        for table in doc.tables:
            rows = table.rows
            i = 0
            while i < len(rows):
                if len(rows[i].cells) < 2:
                    i += 1
                    continue

                row0 = [c.text.strip() for c in rows[i].cells]
                if not row0[0].startswith("QN="):
                    i += 1
                    continue

                content = row0[1]  # Question text
                options = []
                for j in range(1, 5):
                    if i + j >= len(rows): break
                    opt_row = [c.text.strip() for c in rows[i + j].cells]
                    if len(opt_row) >= 2:
                        options.append(f"{opt_row[0]} {opt_row[1]}")
                answer = None
                if i + 5 < len(rows):
                    ans_row = [c.text.strip() for c in rows[i + 5].cells]
                    if ans_row[0].startswith("ANSWER:"):
                        answer = ans_row[1]

                # Save to DB if valid
                if subject and content and len(options) == 4 and answer:
                    q = Question(subject, content, ";".join(options), answer)
                    self.db.add_question(q)
                    print(f"[IMPORTED] {content[:30]}...")

                # Move to next question block
                i += 9  # Assuming fixed row pattern

        return True
