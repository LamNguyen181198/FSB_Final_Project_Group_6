# desktop_app.py
import tkinter as tk
from docx import Document
from tkinter import filedialog as fd, Tk, filedialog
from tkinter import ttk, messagebox
import requests
import os
from PIL import Image, ImageTk
import io
import zipfile
from part2_bacdd1 import db, importer
from part2_bacdd1.models import Question
from part2_bacdd1.db import Database
import random
db = Database()

API_URL = "http://127.0.0.1:8000"

selected_file_path = None
# UI setup
root = tk.Tk()
root.title("Document Management Desktop App")
root.geometry("1200x800")

notebook = ttk.Notebook(root)
tab_doc = tk.Frame(notebook)
tab_question = tk.Frame(notebook)
tab_create_exam = tk.Frame(notebook)
notebook.add(tab_question, text="Exam Document Management")
notebook.add(tab_create_exam, text="Create Exam")
notebook.pack(fill="both", expand=True)

# --- Tab: Question Management ---
tk.Label(tab_question, text="Subject:").pack()
subject_entry = tk.Entry(tab_question, width=50)
subject_entry.pack()

tk.Label(tab_question, text="Question Content:").pack()
content_entry = tk.Entry(tab_question, width=80)
content_entry.pack()

tk.Label(tab_question, text="Options:").pack()
options_frame = tk.Frame(tab_question)
options_frame.pack()

selected_file_label = tk.Label(tab_question, text="No file selected", anchor="w", width=60)
selected_file_label.pack(pady=(5, 0))

file_status_label = tk.Label(tab_question, text="", fg="green")
file_status_label.pack(pady=(5, 0))

option_entries = []
for opt in ["A", "B", "C", "D"]:
    row = tk.Frame(options_frame)
    row.pack(fill="x", pady=2)
    label = tk.Label(row, text=f"Option {opt}:", width=12)
    label.pack(side="left")
    entry = tk.Entry(row, width=60)
    entry.pack(side="left", fill="x", expand=True)
    option_entries.append(entry)

tk.Label(tab_question, text="Correct Answer:").pack()
answer_entry = tk.Entry(tab_question, width=20)
answer_entry.pack()

def add_question():
    subject = subject_entry.get().strip()
    content = content_entry.get().strip()
    answer = answer_entry.get().strip()

    option_labels = ["A", "B", "C", "D"]
    options = []
    for label, entry in zip(option_labels, option_entries):
        text = entry.get().strip()
        if text:
            options.append(f"{label}. {text}")

    options_str = ";".join(options)

    if subject and content and options_str and answer:
        db.add_question(Question(subject, content, options_str, answer))
        refresh_questions()
        messagebox.showinfo("Add Question", "Question added successfully!")

def refresh_questions():
    subject = subject_entry.get().strip()
    question_list.delete(0, tk.END)
    for row in db.get_questions(subject):
        question_list.insert(tk.END, f"{row[0]}: {row[1]} | {row[2]} | Answer: {row[3]}")

def update_question():
    selected = question_list.curselection()
    if selected:
        qid = question_list.get(selected[0]).split(":")[0]
        subject = subject_entry.get().strip()
        content = content_entry.get().strip()
        answer = answer_entry.get().strip()

        option_labels = ["A", "B", "C", "D"]
        options = []
        for label, entry in zip(option_labels, option_entries):
            text = entry.get().strip()
            if text:
                options.append(f"{label}. {text}")

        options_str = ";".join(options)

        db.update_question(qid, subject, content, options_str, answer)
        refresh_questions()
        messagebox.showinfo("Update Question", "Question updated successfully!")

def delete_question():
    selected = question_list.curselection()
    if selected:
        qid = question_list.get(selected[0]).split(":")[0]
        db.delete_question(qid)
        refresh_questions()
        messagebox.showinfo("Delete Question", "Question deleted successfully!")

question_importer = importer.QuestionImporter(db)
def add_and_import_questions():
    global selected_file_path

    path = filedialog.askopenfilename(
        title="Select question file (.docx)",
        filetypes=[("Word files", ".*")]
    )

    if not path:
        messagebox.showwarning("No File", "Please select a .docx file.")
        return

    if not path.lower().endswith(".docx"):
        messagebox.showwarning("Invalid File", "Only .docx files are supported.")
        return

    selected_file_path = path

    try:
        with open(selected_file_path, "rb") as file:
            files = {
                "file": (
                    os.path.basename(selected_file_path),
                    file,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            }
            response = requests.post(f"{API_URL}/add_document", files=files)
            response.raise_for_status()

        question_importer.import_from_docx(selected_file_path)

        messagebox.showinfo("Success", "Document uploaded and questions imported successfully.")

        file_status_label.config(
            text=f"{os.path.basename(path)} added successfully.",
            fg="green"
        )

        selected_file_path = None
        refresh_questions()

    except Exception as e:
        messagebox.showerror("Error", str(e))
        file_status_label.config(
            text="Failed to import file.",
            fg="red"
        )

question_list = tk.Listbox(tab_question, width=100)
question_list.pack(padx=10, pady=10)

question_list.bind('<<ListboxSelect>>', lambda e: on_select_question(e))

def on_select_question(event):
    selected = question_list.curselection()
    if selected:
        qid = question_list.get(selected[0]).split(":")[0]
        subject = subject_entry.get().strip()
        rows = db.get_questions(subject)
        for row in rows:
            if str(row[0]) == qid:
                content_entry.delete(0, tk.END)
                content_entry.insert(0, row[1])

                opts = row[2].split(";")
                for i, e in enumerate(option_entries):
                    e.delete(0, tk.END)
                    if i < len(opts):
                        opt_text = opts[i].strip()
                        if ". " in opt_text:
                            opt_text = opt_text.split(". ", 1)[1]
                        e.insert(0, opt_text)

                answer_entry.delete(0, tk.END)
                answer_entry.insert(0, row[3])
                break

tk.Button(tab_question, text="Import Questions from File", command=add_and_import_questions).pack(pady=10)
tk.Button(tab_question, text="Load Questions", command=refresh_questions).pack()
tk.Button(tab_question, text="Add Question", command=add_question).pack(pady=2)
tk.Button(tab_question, text="Update Question", command=update_question).pack(pady=2)
tk.Button(tab_question, text="Delete Question", command=delete_question).pack(pady=2)

# --- Tab: Create Exam ---
tk.Label(tab_create_exam, text="Subject:").pack()
exam_subject = tk.Entry(tab_create_exam, width=50)
exam_subject.pack()

tk.Label(tab_create_exam, text="Exam Code:").pack()
exam_code = tk.Entry(tab_create_exam, width=80)
exam_code.pack()

tk.Label(tab_create_exam, text="Duration (minutes):").pack()
exam_duration = tk.Entry(tab_create_exam, width=80)
exam_duration.pack()

tk.Label(tab_create_exam, text="Number of Questions:").pack()
num_ques = tk.Entry(tab_create_exam, width=80)
num_ques.pack()

def create_exam():
    subject = exam_subject.get().strip()
    code = exam_code.get().strip()
    duration = exam_duration.get().strip()
    num_questions = num_ques.get().strip()
    question_list = db.get_random_questions(subject, int(num_questions))

    exam_file = Document()
    exam_file.add_heading(f"{subject} Exam", level=1)
    exam_file.add_paragraph(f"Exam Code: {code}; Duration: {duration} minutes")
    index_question = 1
    for question in question_list:
        exam_file.add_paragraph(f"Question {index_question}: {question[1]}")
        index_question += 1
        options = question[2].split(';')
        random.shuffle(options)
        index_option = 1
        for option in options:
            exam_file.add_paragraph(f"{index_option}. {option.split('.')[1]}")
            index_option += 1
    root = Tk()
    root.withdraw()
    file_path = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")],
        title="Save Document As",
        initialfile = f"{subject}_{code}.docx"
    )
    if file_path:
        exam_file.save(file_path)
        print(f"Document saved to: {file_path}")
    else:
        print("Save operation cancelled.")

tk.Button(tab_create_exam, text="Create Exam", command=create_exam).pack(pady=10)

root.mainloop()
